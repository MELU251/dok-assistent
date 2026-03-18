"""Unit tests fuer src/output.py."""

import pytest
from pathlib import Path
from docx import Document as DocxDocument


VALID_SECTIONS = {
    "zusammenfassung": "Kurze Zusammenfassung des Angebots.",
    "technische_loesung": "Beschreibung der technischen Loesung.",
    "lieferumfang": "Liste der Lieferpositionen.",
    "offene_punkte": ["Klaerung Lieferzeit", "Klaerung Zertifizierung"],
}
VALID_SOURCES = ["angebot_2023.pdf", "referenz_projekt.docx"]


class TestWriteDocx:
    def test_creates_file(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        result = write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        assert result.exists()

    def test_returns_resolved_path(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        result = write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        assert result.is_absolute()

    def test_file_is_valid_docx(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        doc = DocxDocument(str(out))
        assert doc is not None

    def test_all_section_headings_present(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        doc = DocxDocument(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Zusammenfassung" in full_text
        assert "Loesung" in full_text
        assert "Lieferumfang" in full_text
        assert "Offene Punkte" in full_text

    def test_sources_appear_in_document(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        doc = DocxDocument(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "angebot_2023.pdf" in full_text
        assert "referenz_projekt.docx" in full_text

    def test_hil_notice_present(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        doc = DocxDocument(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "KI-generierter Entwurf" in full_text

    def test_list_content_rendered_as_bullets(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        write_docx(VALID_SECTIONS, VALID_SOURCES, out)
        doc = DocxDocument(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Klaerung Lieferzeit" in full_text
        assert "Klaerung Zertifizierung" in full_text

    def test_missing_section_raises_value_error(self, tmp_path):
        from src.output import write_docx

        incomplete = {k: v for k, v in VALID_SECTIONS.items() if k != "lieferumfang"}
        with pytest.raises(ValueError, match="lieferumfang"):
            write_docx(incomplete, VALID_SOURCES, tmp_path / "out.docx")

    def test_creates_parent_dirs(self, tmp_path):
        from src.output import write_docx

        nested = tmp_path / "subdir" / "deep" / "out.docx"
        write_docx(VALID_SECTIONS, VALID_SOURCES, nested)
        assert nested.exists()

    def test_empty_sources_handled(self, tmp_path):
        from src.output import write_docx

        out = tmp_path / "test.docx"
        write_docx(VALID_SECTIONS, [], out)
        doc = DocxDocument(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Keine Quelldokumente" in full_text
