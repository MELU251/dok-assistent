"""Ausgabe-Modul: Generiert formatierte .docx-Angebotsentwuerfe.

Nimmt ein strukturiertes Sections-Dict und eine Liste von Quellen entgegen
und schreibt eine formatierte Word-Datei in das angegebene Ausgabeverzeichnis.
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Pflicht-Abschnitte in festgelegter Reihenfolge
SECTION_ORDER = [
    "zusammenfassung",
    "technische_loesung",
    "lieferumfang",
    "offene_punkte",
]

SECTION_LABELS = {
    "zusammenfassung": "1. Zusammenfassung",
    "technische_loesung": "2. Technische Loesung",
    "lieferumfang": "3. Lieferumfang",
    "offene_punkte": "4. Offene Punkte",
}

_HIL_TEXT = (
    "KI-generierter Entwurf — Pflichtpruefung durch Vertriebsingenieur erforderlich "
    "vor Versand an Kunden."
)


def write_docx(
    sections: dict[str, str | list[str]],
    sources: list[str],
    output_path: Path,
) -> Path:
    """Schreibt einen formatierten Angebotsentwurf als .docx-Datei.

    Args:
        sections: Dict mit Pflicht-Keys (zusammenfassung, technische_loesung,
            lieferumfang, offene_punkte). Werte koennen str oder list[str] sein.
        sources: Liste der verwendeten Quelldokumente (Dateinamen).
        output_path: Zielpfad der erzeugten .docx-Datei.

    Returns:
        Absoluter Pfad zur erzeugten Datei.

    Raises:
        ValueError: Wenn ein Pflicht-Abschnitt fehlt.
        OSError: Wenn die Datei nicht geschrieben werden kann.
    """
    _validate_sections(sections)

    doc = DocxDocument()
    _write_header(doc)
    _write_hil_notice(doc)
    _write_sections(doc, sections)
    _write_sources(doc, sources)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Angebotsentwurf geschrieben: %s", output_path)
    return output_path.resolve()


def _validate_sections(sections: dict) -> None:
    """Prueft ob alle Pflicht-Abschnitte vorhanden sind."""
    missing = [key for key in SECTION_ORDER if key not in sections]
    if missing:
        raise ValueError(f"Fehlende Pflicht-Abschnitte: {missing}")


def _write_header(doc: DocxDocument) -> None:
    """Firmen-Platzhalter-Header schreiben."""
    title = doc.add_heading("Angebotsentwurf", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("[Firmenname] — Vertriebsdokument")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # Leerzeile


def _write_hil_notice(doc: DocxDocument) -> None:
    """Human-in-the-Loop-Hinweis als auffaelligen Block schreiben."""
    p = doc.add_paragraph()
    run = p.add_run(_HIL_TEXT)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    doc.add_paragraph()


def _write_sections(doc: DocxDocument, sections: dict) -> None:
    """Alle Pflicht-Abschnitte in festgelegter Reihenfolge schreiben."""
    for key in SECTION_ORDER:
        label = SECTION_LABELS[key]
        content = sections[key]

        doc.add_heading(label, level=1)

        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(content))

        doc.add_paragraph()


def _write_sources(doc: DocxDocument, sources: list[str]) -> None:
    """Quellenangaben-Anhang schreiben."""
    doc.add_heading("Verwendete Quellen", level=2)
    if sources:
        for src in sources:
            doc.add_paragraph(src, style="List Bullet")
    else:
        doc.add_paragraph("Keine Quelldokumente verwendet.")
